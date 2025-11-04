import os
import time
import polars as pl
from docx import Document

# Start the timer
start_time = time.time()

# Define the base and data tables directory
BASE_DIR = os.path.join(os.getcwd(), 'THD Data Warehouse')
DATA_TABLES_DIR = os.path.join(BASE_DIR, 'Reviews and Questions')

# Define the specific Parquet files to summarize
PARQUET_FILES = [
    'online_sales.parquet',
    'online_classification.parquet',
    'online_website_anaylsis.parquet',
    'questions_with_fg_data.parquet',
    'marketing_reviews_to_respond.parquet',
    'questions_without_answer.parquet',
    'FG_processed_data_grouped.parquet',
    'reviews_with_fg_data.parquet'

]

# Create a Word document
doc = Document()
doc.add_heading('Parquet Files Summary', 0)

# Flag to track if any files were processed
files_processed = False

# Process each Parquet file
for file_name in PARQUET_FILES:
    file_path = os.path.join(DATA_TABLES_DIR, file_name)
    
    if os.path.exists(file_path):
        try:
            # Read Parquet file with Polars
            df = pl.read_parquet(file_path)
            total_rows = df.height  # Get total number of rows
            doc.add_heading(file_name, level=1)
            doc.add_paragraph(f"Total Rows: {total_rows:,}")

            # Get sample data (first 2 rows)
            sample_df = df.head(2)

            # Build table in Word
            table_doc = doc.add_table(rows=1, cols=3)
            table_doc.style = 'Table Grid'
            hdr = table_doc.rows[0].cells
            hdr[0].text = 'Column Name'
            hdr[1].text = 'Sample Values'
            hdr[2].text = 'Data Type'

            # Iterate through columns
            for col in df.columns:
                # Get sample values, handling nulls and categoricals
                sample_vals = sample_df[col].to_list()
                sample_str = ', '.join(str(val) if val is not None else 'null' for val in sample_vals)
                # Get Polars data type
                dtype = str(df.schema[col])
                
                row = table_doc.add_row().cells
                row[0].text = col
                row[1].text = sample_str
                row[2].text = dtype

            files_processed = True

        except Exception as e:
            doc.add_paragraph(f"Failed to read {file_name}: {str(e)}")
    else:
        doc.add_paragraph(f"File not found: {file_name}")

# If no files were processed, add a message
if not files_processed:
    doc.add_paragraph("No Parquet files were found or processed successfully.")

# Save the Word document
output_path = os.path.join(BASE_DIR, 'Reviews and Questions', 'parquet_summary_reviews.docx')
doc.save(output_path)

# End the timer and print elapsed time
end_time = time.time()
elapsed_time = end_time - start_time
print(f"Summary saved to {output_path}")
print(f"Script completed in {elapsed_time:.2f} seconds.")