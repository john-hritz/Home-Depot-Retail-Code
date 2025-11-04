import os
import time
import polars as pl
from docx import Document
import re

# Start the timer
start_time = time.time()

# Define the base and data tables directory
BASE_DIR = os.path.join(os.getcwd(), 'THD Data Warehouse')
DATA_TABLES_DIR = os.path.join(BASE_DIR, 'Data_Tables')

# Define the specific Parquet files created by the previous script
PARQUET_FILES = [
    'online_sales.parquet',
    'online_classification.parquet',
    'online_website_anaylsis.parquet',
    'BA_scorecard.parquet',
    'online_stores.parquet',
    'calendar.parquet',
    'fg_status.parquet',
    'merged_classification.parquet'
]

# Function to parse and sort week values
def parse_and_sort_weeks(week_values):
    # Regular expression to extract week and year (e.g., "Fiscal Week 17 of 2025")
    pattern = r'Fiscal Week (\d+) of (\d{4})'
    
    # List to store parsed (year, week, original_string) tuples
    parsed_weeks = []
    
    for val in week_values:
        if val is None:
            continue
        match = re.match(pattern, str(val))
        if match:
            week = int(match.group(1))
            year = int(match.group(2))
            parsed_weeks.append((year, week, val))
        else:
            # Handle malformed entries by keeping them as-is (will sort to end)
            parsed_weeks.append((0, 0, val))
    
    # Sort in descending order: first by year, then by week
    sorted_weeks = sorted(parsed_weeks, key=lambda x: (x[0], x[1]), reverse=True)
    
    # Return the original string values in sorted order
    return [week[2] for week in sorted_weeks]

# Create a Word document
doc = Document()
doc.add_heading('Parquet Files Summary', 0)

# Flag to track if any files were processed
files_processed = False

# Process each Parquet file
for file_name in PARQUET_FILES:
    file_path = os.path.join(DATA_TABLES_DIR, file_name)
    
    if os.path.exists(file_path):
        try:
            # Read Parquet file with Polars
            df = pl.read_parquet(file_path)
            total_rows = df.height  # Get total number of rows
            doc.add_heading(file_name, level=1)
            doc.add_paragraph(f"Total Rows: {total_rows:,}")

            # Check for unique values of 'week' column for specific files
            if file_name in ['online_sales.parquet', 'online_website_anaylsis.parquet']:
                if 'week' in df.columns:
                    unique_weeks = df['week'].drop_nulls().unique().to_list()
                    sorted_weeks = parse_and_sort_weeks(unique_weeks)
                    unique_weeks_str = ', '.join(str(val) for val in sorted_weeks)
                    doc.add_paragraph(f"Unique Week Values (Sorted High to Low): {unique_weeks_str if sorted_weeks else 'None'}")
                else:
                    doc.add_paragraph("Unique Week Values: Column 'week' not found")

            # Get schema
            schema = df.schema
            doc.add_paragraph("\nSchema:")
            schema_table = doc.add_table(rows=1, cols=2)
            schema_table.style = 'Table Grid'
            hdr = schema_table.rows[0].cells
            hdr[0].text = 'Column Name'
            hdr[1].text = 'Data Type'

            for col, dtype in schema.items():
                row = schema_table.add_row().cells
                row[0].text = col
                row[1].text = str(dtype)

            # Get sample data (first 2 rows)
            doc.add_paragraph("\nSample Data (First 2 Rows):")
            sample_df = df.head(2)
            sample_table = doc.add_table(rows=1, cols=2)
            sample_table.style = 'Table Grid'
            hdr = sample_table.rows[0].cells
            hdr[0].text = 'Column Name'
            hdr[1].text = 'Sample Values'

            for col in df.columns:
                sample_vals = sample_df[col].to_list()
                sample_str = ', '.join(str(val) if val is not None else 'null' for val in sample_vals)
                row = sample_table.add_row().cells
                row[0].text = col
                row[1].text = sample_str

            # Get basic statistics
            doc.add_paragraph("\nColumn Statistics:")
            stats_table = doc.add_table(rows=1, cols=4)
            stats_table.style = 'Table Grid'
            hdr = stats_table.rows[0].cells
            hdr[0].text = 'Column Name'
            hdr[1].text = 'Null Count'
            hdr[2].text = 'Unique Values'
            hdr[3].text = 'Min/Max (Numeric) or Example Values (Categorical/String)'

            for col in df.columns:
                null_count = df[col].null_count()
                unique_count = df[col].n_unique()
                dtype = schema[col]
                
                # Determine min/max or example values based on data type
                if dtype in [pl.Int32, pl.Int64, pl.Float32, pl.Float64]:
                    min_val = df[col].min()
                    max_val = df[col].max()
                    stat_str = f"Min: {min_val if min_val is not None else 'null'}, Max: {max_val if max_val is not None else 'null'}"
                elif dtype in [pl.Categorical, pl.Utf8]:
                    example_vals = df[col].drop_nulls().head(3).to_list()
                    stat_str = ', '.join(str(val) if val is not None else 'null' for val in example_vals)
                else:
                    stat_str = 'N/A'

                row = stats_table.add_row().cells
                row[0].text = col
                row[1].text = str(null_count)
                row[2].text = str(unique_count)
                row[3].text = stat_str

            files_processed = True

        except Exception as e:
            doc.add_paragraph(f"Failed to read {file_name}: {str(e)}")
    else:
        doc.add_paragraph(f"File not found: {file_name}")

# If no files were processed, add a message
if not files_processed:
    doc.add_paragraph("No Parquet files were found or processed successfully.")

# Save the Word document
output_path = os.path.join(DATA_TABLES_DIR, 'parquet_summary.docx')
doc.save(output_path)

# End the timer and print elapsed time
end_time = time.time()
elapsed_time = end_time - start_time
print(f"Summary saved to {output_path}")
print(f"Script completed in {elapsed_time:.2f} seconds.")